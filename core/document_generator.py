#== == == == == == == == == == core / document_generator.py == == == == == == == == == ==
# !/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档生成器 - 支持 Word 和 PDF 格式
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re
from datetime import datetime


class DocumentGenerator:
    """教学文档生成器"""

    def __init__(self):
        self.doc = None

    def save_document(self, content: str, data: dict, save_path: str):
        """
        保存文档

        Args:
            content: AI生成的内容
            data: 表单数据
            save_path: 保存路径
        """
        ext = Path(save_path).suffix.lower()

        if ext == ".docx":
            self._save_word(content, data, save_path)
        elif ext == ".pdf":
            # 先生成Word再转换（需要额外依赖）
            temp_word = save_path.replace(".pdf", ".docx")
            self._save_word(content, data, temp_word)
            self._convert_to_pdf(temp_word, save_path)
        elif ext == ".txt":
            self._save_txt(content, save_path)
        else:
            # 默认保存为Word
            if not save_path.endswith(".docx"):
                save_path += ".docx"
            self._save_word(content, data, save_path)

    def _save_word(self, content: str, data: dict, save_path: str):
        """保存为Word文档"""
        self.doc = Document()

        # 添加标题
        title = self.doc.add_heading(f"{data['course_name']} - {data['doc_type']}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加基本信息
        self.doc.add_heading("基本信息", level=1)
        info_table = self.doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'

        info_data = [
            ("课程名称", data['course_name']),
            ("教学专业", data['major']),
            ("授课对象", data['audience']),
            ("课时安排", f"{data['hours']} 学时"),
            ("教学模式", data['teaching_mode']),
            ("考核方式", data['assessment'])
        ]

        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value

        # 添加AI生成内容
        self.doc.add_heading("文档内容", level=1)

        # 解析Markdown内容并添加
        self._parse_markdown_to_docx(content)

        # 添加页脚
        section = self.doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | 教学文档智能生成系统"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 保存
        self.doc.save(save_path)

    def _parse_markdown_to_docx(self, content: str):
        """解析Markdown内容到Word"""
        lines = content.split('\n')
        current_para = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 标题处理
            if line.startswith('# '):
                self.doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                self.doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                self.doc.add_heading(line[4:], level=3)
            # 列表处理
            elif line.startswith('- ') or line.startswith('* '):
                p = self.doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                p = self.doc.add_paragraph(line, style='List Number')
            # 普通段落
            else:
                self.doc.add_paragraph(line)

    def _save_txt(self, content: str, save_path: str):
        """保存为文本文件"""
        with open(save_path, 'w', encoding='utf-8') as f:
            f.write(content)

    def _convert_to_pdf(self, word_path: str, pdf_path: str):
        """Word转PDF（需要libreoffice或docx2pdf）"""
        try:
            # 尝试使用docx2pdf（Windows）
            from docx2pdf import convert
            convert(word_path, pdf_path)
        except ImportError:
            # 降级处理：复制Word文件并重命名提示
            import shutil
            shutil.copy(word_path, pdf_path.replace('.pdf', '.docx'))
            raise Warning("PDF转换需要安装 docx2pdf 或 LibreOffice，已保存为Word格式")